
import os
import logging
from typing import Dict, Optional, List, Union

from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader
from docx import Document
from config.get_embedding import get_embedding_gemini
from models.knowledge_base import DocumentChunk
from config.database import AsyncSessionLocal
from sqlalchemy import delete, select
from bs4 import BeautifulSoup # Cần cho rich text: pip install beautifulsoup4
from config.chromadb_config import add_documents_to_chroma, delete_documents_by_detail_id

logger = logging.getLogger(__name__)


async def insert_chunks_to_db(chunks_data: list):
    try:
        await add_documents_to_chroma(chunks_data)
    except Exception as e:
        logger.error(f"Lỗi bulk insert chunks: {str(e)}")
        raise
        


async def delete_chunks_by_detail_id(detail_id: int):
    """
    Xóa tất cả chunks liên quan đến một knowledge_base_detail_id
    Từ cả PostgreSQL VÀ ChromaDB
    
    Args:
        detail_id: ID của knowledge_base_detail
    """
    try:
        # Xóa từ PostgreSQL
        # await session.execute(
        #     delete(DocumentChunk).where(DocumentChunk.knowledge_base_detail_id == detail_id)
        # )
        # await session.commit()
        
        # Xóa từ ChromaDB
        await delete_documents_by_detail_id(detail_id)
        
        logger.info(f"Đã xóa tất cả chunks của detail_id={detail_id} từ PostgreSQL và ChromaDB")
        return True
    except Exception as e:
        logger.error(f"Lỗi xóa chunks: {str(e)}")
        return False
        


class FileProcessor:
    """Class xử lý đọc nội dung từ các loại file khác nhau sử dụng Langchain loaders"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.xlsx', '.xls'}
    
    # Cấu hình chunking
    CHUNK_SIZE = 1500
    CHUNK_OVERLAP = 200
    
    @staticmethod
    def is_supported_file(filename: str) -> bool:
        """Kiểm tra xem file có được hỗ trợ không"""
        ext = os.path.splitext(filename)[1].lower()
        return ext in FileProcessor.SUPPORTED_EXTENSIONS
    
    @staticmethod
    async def extract_text_from_pdf(file_path: str) -> Dict[str, any]:
        """
        Đọc nội dung từ file PDF sử dụng pypdf
        """
        try:
            reader = PdfReader(file_path)
            
            # Lấy nội dung từ tất cả các trang
            full_text = ""
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n\n"
            
            # Metadata
            metadata = {
                'num_pages': len(reader.pages),
                'file_type': 'PDF',
                'source': file_path
            }
            
            # Thêm metadata từ PDF nếu có
            if reader.metadata:
                metadata.update({
                    'title': reader.metadata.get('/Title', ''),
                    'author': reader.metadata.get('/Author', ''),
                    'subject': reader.metadata.get('/Subject', ''),
                })
            
            return {
                'success': True,
                'content': full_text.strip(),
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Lỗi đọc file PDF: {str(e)}")
            return {
                'success': False,
                'content': '',
                'error': str(e)
            }
    
    @staticmethod
    async def extract_text_from_docx(file_path: str) -> Dict[str, any]:
        """
        Đọc nội dung từ file DOCX sử dụng python-docx
        """
        try:
            doc = Document(file_path)
            text_content = []
            
            # Đọc các đoạn văn
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Đọc nội dung trong bảng
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_content.append(row_text)
            
            full_text = "\n\n".join(text_content)
            
            # Metadata
            metadata = {
                'num_paragraphs': len(doc.paragraphs),
                'num_tables': len(doc.tables),
                'file_type': 'DOCX',
                'source': file_path
            }
            
            # Thêm core properties nếu có
            if doc.core_properties:
                metadata.update({
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'subject': doc.core_properties.subject or '',
                })
            
            return {
                'success': True,
                'content': full_text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Lỗi đọc file DOCX: {str(e)}")
            return {
                'success': False,
                'content': '',
                'error': str(e)
            }
    
    @staticmethod
    async def extract_text_from_excel(file_path: str) -> Dict[str, any]:
        """
        Đọc nội dung từ file Excel sử dụng UnstructuredExcelLoader
        """
        try:
            import pandas as pd
            
            # Đọc tất cả các sheet
            excel_file = pd.ExcelFile(file_path)
            sheet_contents = []
            total_rows = 0
            
            for sheet_name in excel_file.sheet_names:
                try:
                    # Đọc sheet
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    if df.empty:
                        continue
                    
                    # Chuyển DataFrame thành text có cấu trúc
                    # Format: Sheet name, header, và từng dòng dữ liệu
                    sheet_text = f"=== SHEET: {sheet_name} ===\n\n"
                    
                    # Thêm header
                    headers = df.columns.tolist()
                    sheet_text += "HEADER: " + " | ".join([str(h) for h in headers]) + "\n\n"
                    
                    # Thêm từng dòng dữ liệu với header
                    for idx, row in df.iterrows():
                        row_text = []
                        for col in df.columns:
                            value = row[col]
                            # Bỏ qua giá trị NaN
                            if pd.notna(value):
                                row_text.append(f"{col}: {value}")
                        
                        if row_text:
                            sheet_text += " | ".join(row_text) + "\n"
                    
                    sheet_contents.append(sheet_text)
                    total_rows += len(df)
                    
                except Exception as e:
                    logger.warning(f"Lỗi đọc sheet '{sheet_name}': {str(e)}")
                    continue
            
            if not sheet_contents:
                return {
                    'success': False,
                    'content': '',
                    'error': 'Không thể đọc nội dung từ file Excel'
                }
            
            # Kết hợp tất cả các sheet
            full_text = "\n\n".join(sheet_contents)
            
            # Metadata
            metadata = {
                'num_sheets': len(excel_file.sheet_names),
                'total_rows': total_rows,
                'file_type': 'Excel',
                'source': file_path
            }
            
            return {
                'success': True,
                'content': full_text,
                'metadata': metadata,
                'sheet_contents': sheet_contents  # Để chunk theo sheet nếu cần
            }
            
        except Exception as e:
            logger.error(f"Lỗi đọc file Excel: {str(e)}")
            return {
                'success': False,
                'content': '',
                'error': str(e)
            }
    
    @classmethod
    async def process_file(cls, file_path: str, filename: str) -> Dict[str, any]:
        """
        Xử lý file dựa trên extension
        """
        ext = os.path.splitext(filename)[1].lower()
        
        if ext == '.pdf':
            result = await cls.extract_text_from_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            result = await cls.extract_text_from_docx(file_path)
        elif ext in ['.xlsx', '.xls']:
            result = await cls.extract_text_from_excel(file_path)
        else:
            return {
                'success': False,
                'content': '',
                'error': f'Định dạng file không được hỗ trợ: {ext}'
            }
        
        # Thêm tên file vào metadata
        if result.get('success'):
            result['metadata']['filename'] = filename
            result['metadata']['file_extension'] = ext
        
        return result
    
    @classmethod
    async def process_and_chunk_file(
        cls, 
        file_path: str, 
        filename: str,
        knowledge_base_detail_id: int
    ) -> Dict[str, any]:
        """
        Xử lý file, chunk nội dung và lưu vào database
        """
        try:
            # Bước 1: Extract text từ file
            result = await cls.process_file(file_path, filename)
            
            if not result.get('success'):
                return result
            
            # Bước 2: Chunk nội dung
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=cls.CHUNK_SIZE,
                chunk_overlap=cls.CHUNK_OVERLAP
            )
            
            # Split từ text content
            all_chunks = text_splitter.split_text(result['content'])
            
            if not all_chunks:
                return {
                    'success': False,
                    'error': 'Không có nội dung để chunk'
                }
            
            # === SỬA ĐỔI: BATCH EMBEDDING ===
            # 1. Gom tất cả text vào một list
            all_texts = [chunk for chunk in all_chunks if chunk] 
            logger.info(f"Chuẩn bị tạo embedding cho {len(all_texts)} chunks từ file {filename}.")

            # 2. Gọi API embedding MỘT LẦN cho cả batch
            try:
                all_vectors = await get_embedding_gemini(all_texts)
                logger.info(f"Đã tạo thành công {len(all_vectors)} vectors.")
            except Exception as e:
                logger.error(f"Lỗi khi gọi batch embedding: {str(e)}")
                return {'success': False, 'error': f"Lỗi embedding: {str(e)}"}

            # 3. Chuẩn bị data
            chunks_data = []
            for text, vector in zip(all_texts, all_vectors):
                chunks_data.append({
                    'chunk_text': text,
                    'search_vector': vector, # Giả sử hàm gemini trả về list
                    'knowledge_base_detail_id': knowledge_base_detail_id
                })
            
            # 4. Insert MỘT LẦN vào CSDL (dùng hàm đã sửa)
            if chunks_data:
                await insert_chunks_to_db(chunks_data)
            
            return {
                'success': True,
                'message': f'Đã xử lý file {filename}',
                'chunks_created': len(chunks_data),
                'metadata': result['metadata']
            }
            # === KẾT THÚC SỬA ĐỔI ===
            
        except Exception as e:
            logger.error(f"Lỗi xử lý và chunk file {filename}: {str(e)}")
            return {
                'success': False,
                'error': str(e)
            }


async def process_uploaded_file(
    file_path: str, 
    filename: str,
    knowledge_base_detail_id: int = None
) -> Dict[str, any]:
    """
    Function wrapper để xử lý file upload
    """
    if knowledge_base_detail_id is not None:
        # Xử lý và chunk vào database
        return await FileProcessor.process_and_chunk_file(
            file_path, 
            filename, 
            knowledge_base_detail_id
        )
    else:
        # Chỉ extract text, không chunk
        return await FileProcessor.process_file(file_path, filename)


# ==========================================================
# HÀM MỚI ĐỂ XỬ LÝ RICH TEXT (NHẬP THỦ CÔNG)
# ==========================================================
async def process_rich_text(
    raw_content: str, 
    knowledge_base_detail_id: int
) -> Dict[str, any]:
    """
    Xử lý nội dung rich text (HTML), chunk, tạo vector và lưu vào CSDL
    """
    try:
        # Bước 1: Làm sạch HTML để lấy text
        soup = BeautifulSoup(raw_content, "html.parser")
        text_content = soup.get_text(separator="\n", strip=True)
        
        if not text_content:
            return {'success': False, 'error': 'Nội dung text rỗng sau khi lọc HTML'}

        # Bước 2: Chunk text
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=FileProcessor.CHUNK_SIZE,
            chunk_overlap=FileProcessor.CHUNK_OVERLAP
        )
        all_chunks = text_splitter.split_text(text_content)
        
        if not all_chunks:
            return {'success': False, 'error': 'Không có nội dung để chunk sau khi lọc'}

        logger.info(f"Đã chunk nội dung rich text thành {len(all_chunks)} chunks.")

        # Bước 3: Tạo Embeddings (Batch)
        all_texts = [chunk for chunk in all_chunks if chunk]
        
        try:
            all_vectors = await get_embedding_gemini(all_texts)
            logger.info(f"Đã tạo thành công {len(all_vectors)} vectors cho rich text.")
        except Exception as e:
            logger.error(f"Lỗi khi gọi batch embedding cho rich text: {str(e)}")
            return {'success': False, 'error': f"Lỗi embedding: {str(e)}"}
        
        # Bước 4: Chuẩn bị data và Lưu vào CSDL (Batch)
        chunks_data = []
        for text, vector in zip(all_texts, all_vectors):
            chunks_data.append({
                'chunk_text': text,
                'search_vector': vector,
                'knowledge_base_detail_id': knowledge_base_detail_id
            })
        
        if chunks_data:
            await insert_chunks_to_db(chunks_data) # Gọi hàm đã sửa
        
        return {
            'success': True,
            'chunks_created': len(chunks_data),
            'source_type': 'RICH_TEXT'
        }

    except Exception as e:
        logger.error(f"Lỗi khi xử lý rich text: {str(e)}")
        return {'success': False, 'error': str(e)}