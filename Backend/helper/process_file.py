import logging
from typing import Any, Dict, List, Optional
from datetime import datetime, date, time
from pypdf import PdfReader
from docx import Document
import pandas as pd
import json

logger = logging.getLogger(__name__)



async def extract_text_from_pdf(file_path: str) -> Optional[str]:
    try:
        reader = PdfReader(file_path)
        contents = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                contents.append(text.strip())
                

        return "\n\n".join(contents).strip()

    except Exception as e:
        logger.error(f"Lỗi đọc file PDF: {e}")
        return None

async def extract_text_from_docx(file_path: str) -> Optional[str]:

    try:
        doc = Document(file_path)
        contents = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                contents.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    contents.append(row_text)

        
        
        return "\n\n".join(contents).strip()

    except Exception as e:
        return None

async def extract_text_from_excel(file_path: str) -> Optional[str]:

    sheet_jsons = {}

    try:
        with pd.ExcelFile(file_path) as excel:
            for sheet in excel.sheet_names:
                try:
                    df = pd.read_excel(excel, sheet_name=sheet)
                    
                    df = df.map(lambda x: x.isoformat() if isinstance(x, (pd.Timestamp, datetime, date, time)) else x)
                    if df.empty:
                        continue
                    
                    rows_list = []
                    for idx, row in df.iterrows():
                        row_dict = {col: val for col, val in row.items() if pd.notna(val)}
                        if row_dict:
                            rows_list.append(row_dict)

                    if rows_list:
                        sheet_jsons[sheet] = rows_list

                except Exception as e:
                    logger.error(f"Lỗi xử lý sheet '{sheet}': {e}")
                    continue
        
        if not sheet_jsons:
            return None
        
        final_str = "{\n"
        for sheet_name, rows in sheet_jsons.items():
            final_str += f'  "{sheet_name}": [\n'

            for row in rows:
                row_parts = []
                for k, v in row.items():
                    v_str = f'"{v}"' if isinstance(v, str) else str(v)
                    row_parts.append(f'"{k}": {v_str}')
                
                row_str = "{ " + ", ".join(row_parts) + " }"
                final_str += f"    {row_str},\n"

            final_str += "  ],\n"

        final_str += "}"
        
        return final_str
    
    except Exception as e:
        logger.error(f"Lỗi đọc file Excel {file_path}: {e}", exc_info=True)
        return None



async def extract_procedures_from_excel_tthc(file_path: str) -> List[Dict[str, Any]]:
    results = []

    try:
        with pd.ExcelFile(file_path) as excel:
            for sheet in excel.sheet_names:
                try:
                    df = pd.read_excel(excel, sheet_name=sheet)

                    if df.empty:
                        continue
                        
                    if "Tên thủ tục" not in df.columns:
                        continue

                    for _, row in df.iterrows():
                        if pd.isna(row["Tên thủ tục"]):
                            continue

                        procedure_name = str(row["Tên thủ tục"]).strip()

                        metadata_json = {}
                        for col, val in row.items():
                            if pd.notna(val):
                                metadata_json[col] = val

                        results.append({
                            "procedure_name": procedure_name,
                            "metadata_json": metadata_json
                        })
                        
                except Exception as e:
                    logger.error(f"Lỗi xử lý sheet '{sheet}': {e}")
                    continue

        return results
    
    except Exception as e:
        logger.error(f"Lỗi đọc file Excel TTHC {file_path}: {e}", exc_info=True)
        return []

